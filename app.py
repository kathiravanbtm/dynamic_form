from flask import Flask, render_template, request, flash, redirect, url_for, send_from_directory
from werkzeug.utils import secure_filename
import os
import re
from docx import Document

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"docx"}

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_placeholders(docx_path):
    """Extracts placeholders as single inputs or loop structures."""
    doc = Document(docx_path)

    text_inputs = set()  # Stores single-line placeholders
    loops = {}  # Stores loop structures

    # Regex patterns for easy-template-x syntax
    loop_open_regex = r"{#(\w+)}"  # Match {#loop_name}
    loop_close_regex = r"{/}"  # Match {/}
    variable_regex = r"{(\w+)}"  # Match {variable}

    in_loop = False
    loop_name = None

    for para in doc.paragraphs:
        text = para.text.strip()

        match_loop_open = re.search(loop_open_regex, text)
        match_loop_close = re.search(loop_close_regex, text)

        if match_loop_open:
            in_loop = True
            loop_name = match_loop_open.group(1)
            loops[loop_name] = set()
            continue

        if match_loop_close:
            in_loop = False
            loop_name = None
            continue

        if in_loop and loop_name:
            fields = re.findall(variable_regex, text)
            loops[loop_name].update(fields)
            continue

        # Handle standalone text placeholders
        placeholders = re.findall(variable_regex, text)
        text_inputs.update(placeholders)

    return {
        "single": list(text_inputs),
        "loops": {key: list(value) for key, value in loops.items()},
    }



def extract_placeholders(docx_path):
    """Extracts placeholders, including loops and single inputs."""
    doc = Document(docx_path)

    text_inputs = set()  # Stores single-line placeholders
    loops = {}  # Stores loop structures

    # Regex patterns with flexible spacing
    loop_open_regex = r"{#\s*(\w+)\s*}"  # Match {#loop_name}
    loop_close_regex = r"{/\s*(\w+)\s*}"  # Match {/loop_name}
    variable_regex = r"{\s*(\w+)\s*}"  # Match {variable}

    in_loop = False
    loop_name = None

    for para in doc.paragraphs:
        text = para.text.strip()

        # Detect loop start
        match_loop_open = re.search(loop_open_regex, text)
        if match_loop_open:
            in_loop = True
            loop_name = match_loop_open.group(1)
            loops.setdefault(loop_name, set())
            continue

        # Detect loop end
        match_loop_close = re.search(loop_close_regex, text)
        if match_loop_close:
            in_loop = False
            loop_name = None
            continue

        # Extract placeholders
        placeholders = re.findall(variable_regex, text)

        if in_loop and loop_name:
            loops[loop_name].update(placeholders)
        else:
            text_inputs.update(placeholders)

    return {
        "single": list(text_inputs),
        "loops": {key: list(value) for key, value in loops.items()},
    }

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if "file" not in request.files:
            flash("No file part")
            return redirect(request.url)

        file = request.files["file"]
        if file.filename == "":
            flash("No selected file")
            return redirect(request.url)

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(file_path)
            return redirect(url_for("form", filename=filename))

    return render_template("upload.html")


@app.route("/form")
def form():
    print("hello")
    filename = request.args.get("filename")
    if not filename:
        flash("No file selected")
        return redirect(url_for("index"))

    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(file_path):
        flash("File not found")
        return redirect(url_for("index"))
    print("hello")
    placeholders = extract_placeholders(file_path)
    print("Extracted placeholders:", placeholders)
    return render_template("form.html", filename=filename, placeholders=placeholders)


@app.route("/uploads/<filename>")
def uploaded_file(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)


if __name__ == "__main__":
    app.run(debug=True)
